"""Word 导出（LegalDocGen 思路）— 使用 Word 内置标题样式，避免手搓字体导致版式怪异。

LegalDocGen 的 export_service 核心做法：
- doc.add_heading() 映射 Markdown # / ## / ###
- 正文 add_paragraph()，不强制方正小标宋/仿宋等（由 Word 主题渲染）
- 不用复杂 XML 水印、章节前分页

本模块供 LaborAid 默认导出使用；法院精细排版见 word_export.py（court 模式）。
"""

from __future__ import annotations

import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.docgen.structured.helpers import strip_no_indent_marker
from app.services.docgen.word_export import _SAFE_FILENAME_RE, _sanitize_xml_text


# 中文大写数字映射
_CN_UPPER_DIGITS = "零壹贰叁肆伍陆柒捌玖"
_CN_UPPER_UNITS = ["", "拾", "佰", "仟"]
_CN_UPPER_BIG_UNITS = ["", "万", "亿", "兆"]

# 金额匹配正则：匹配 ¥100,000.00 或 ￥100000 或 人民币10000元 等格式
_AMOUNT_PATTERN = re.compile(
    r"(?:人民币|￥|¥)\s*(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)\s*元?"
)


def _num_to_cn_upper(amount_str: str) -> str:
    """将阿拉伯数字金额转为中文大写金额（如 100000 → 壹拾万元整）。"""
    s = amount_str.replace(",", "").strip()
    if not s:
        return ""
    try:
        num = float(s)
    except ValueError:
        return amount_str

    integer_part = int(num)
    decimal_part = round((num - integer_part) * 100)
    jiao = decimal_part // 10
    fen = decimal_part % 10

    if integer_part == 0:
        result = "零"
    else:
        s_int = str(integer_part)
        n = len(s_int)
        result_parts = []
        zero_flag = False
        for i, ch in enumerate(s_int):
            digit = int(ch)
            pos = n - 1 - i
            unit_idx = pos % 4
            big_idx = pos // 4
            if digit == 0:
                zero_flag = True
                if unit_idx == 0 and big_idx > 0:
                    result_parts.append(_CN_UPPER_BIG_UNITS[big_idx])
                    zero_flag = False
            else:
                if zero_flag:
                    result_parts.append("零")
                    zero_flag = False
                result_parts.append(_CN_UPPER_DIGITS[digit] + _CN_UPPER_UNITS[unit_idx])
                if unit_idx == 0 and big_idx > 0:
                    result_parts.append(_CN_UPPER_BIG_UNITS[big_idx])
        result = "".join(result_parts)

    if jiao == 0 and fen == 0:
        result += "元整"
    elif jiao == 0 and fen > 0:
        result += "元零" + _CN_UPPER_DIGITS[fen] + "分"
    elif jiao > 0 and fen == 0:
        result += "元" + _CN_UPPER_DIGITS[jiao] + "角整"
    else:
        result += "元" + _CN_UPPER_DIGITS[jiao] + "角" + _CN_UPPER_DIGITS[fen] + "分"

    return "人民币" + result


def _auto_uppercase_amounts(text: str) -> str:
    """在文本中检测 ¥/￥/人民币 金额，自动补充中文大写。

    例：¥100,000.00 → ¥100,000.00（人民币壹拾万元整）
    已有大写的不再重复添加。
    """
    def _replacer(m: re.Match) -> str:
        original = m.group(0)
        amount_str = m.group(1)
        cn_upper = _num_to_cn_upper(amount_str)
        return f"{original}（{cn_upper}）"

    # 如果已有中文大写金额，不再转换
    if re.search(r"[\u58f9\u8d30\u53c1\u8086\u4f0d\u9646\u67d2\u634c\u7396\u62fe\u4f70\u4edf\u4e07\u4ebf]", text):
        return text

    return _AMOUNT_PATTERN.sub(_replacer, text)

# 正文字体：对齐法院标准排版（与 word_export.py COURT_FONT_SETTINGS 一致）
_BODY_FONT = "仿宋"
_BODY_FONT_FALLBACK = "FangSong"
_BODY_SIZE = Pt(16)           # 三号 = 16pt（法院标准）
_LINE_SPACING = Pt(28.8)      # 固定值 28.8 磅
_FIRST_LINE_INDENT = Cm(0.74) # 2 字符 ≈ 0.74cm（三号字）

# 标题字体
_TITLE_FONT = "方正小标宋简体"
_TITLE_FONT_ALT = "FZXiaoBiaoSong-B05S"
_HEADING1_FONT = "黑体"
_HEADING1_FONT_FALLBACK = "SimHei"
_HEADING2_FONT = "楷体"
_HEADING2_FONT_FALLBACK = "KaiTi"

# 英文字体
_EN_FONT = "Times New Roman"


def _set_east_asia_font(run, name: str) -> None:
    """设置东亚字体（同时设置拉丁字体名和 w:eastAsia 属性）。"""
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _style_set_east_asia(style, latin: str, east_asia: str) -> None:
    style.font.name = latin
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)


def _configure_document_styles(doc: Document) -> None:
    """统一 Normal / Heading 样式，打开 Word 时版式稳定。"""
    normal = doc.styles["Normal"]
    _style_set_east_asia(normal, _EN_FONT, _BODY_FONT)
    normal.font.size = _BODY_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    heading_configs = [
        (1, 22, _TITLE_FONT_ALT, _TITLE_FONT, False),
        (2, 16, _HEADING1_FONT, _HEADING1_FONT, True),
        (3, 16, _HEADING2_FONT, _HEADING2_FONT, True),
        (4, 14, _EN_FONT, _BODY_FONT, True),
    ]
    for level, size, latin_font, ea_font, bold in heading_configs:
        style_name = f"Heading {level}"
        if style_name not in doc.styles:
            continue
        h = doc.styles[style_name]
        _style_set_east_asia(h, latin_font, ea_font)
        h.font.bold = bold
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(0, 0, 0)
        h_pf = h.paragraph_format
        h_pf.space_before = Pt(6)
        h_pf.space_after = Pt(3)
        if level == 1:
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _setup_page_simple(doc: Document) -> None:
    """A4 纸张，法院标准页边距。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def _add_page_number(doc: Document) -> None:
    """在页脚添加居中页码（- X - 格式）。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        run_pre = p.add_run("- ")
        _set_east_asia_font(run_pre, _BODY_FONT)
        run_pre.font.size = Pt(10)

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_field = p.add_run()
        run_field._element.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run_instr = p.add_run()
        run_instr._element.append(instr_text)
        _set_east_asia_font(run_instr, _BODY_FONT)
        run_instr.font.size = Pt(10)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end = p.add_run()
        run_end._element.append(fld_char_end)

        run_suf = p.add_run(" -")
        _set_east_asia_font(run_suf, _BODY_FONT)
        run_suf.font.size = Pt(10)


def _parse_inline_to_paragraph(p, text: str, *, bold_labels: bool = True) -> None:
    """支持 **加粗**；**标签**：值 常见于法律文书。自动转换金额为大写。"""
    safe = _sanitize_xml_text(text)
    if not safe:
        return
    safe = _auto_uppercase_amounts(safe)
    parts = safe.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        _set_east_asia_font(run, _BODY_FONT)
        run.font.name = _EN_FONT
        run.font.size = _BODY_SIZE
        if i % 2 == 1:
            run.bold = True
        elif bold_labels and (part.rstrip().endswith("：") or part.rstrip().endswith(":")):
            run.bold = True


def _add_body_paragraph(doc: Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = _FIRST_LINE_INDENT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = _LINE_SPACING
    p.paragraph_format.space_after = Pt(4)
    _parse_inline_to_paragraph(p, text)


def add_markdown_to_document(
    doc: Document,
    content: str,
    *,
    first_line_indent: bool = True,
) -> None:
    """将 Markdown 写入 Document（LegalDocGen 风格 + 加粗行内）。"""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        stripped, no_indent = strip_no_indent_marker(stripped)
        if no_indent:
            _add_body_paragraph(doc, stripped, indent=False)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if re.match(r"^\|[\s\-:|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols)
                table.style = "Table Grid"
                table.autofit = True
                for ri, row in enumerate(table_rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(24)
                        run = p.add_run(_sanitize_xml_text(cell_text))
                        _set_east_asia_font(run, _BODY_FONT)
                        run.font.size = Pt(14)
                        run.font.name = _EN_FONT
                        if ri == 0:
                            run.bold = True
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:fill"), "D9E2F3")
                            shading.set(qn("w:val"), "clear")
                            cell._element.get_or_add_tcPr().append(shading)
            continue

        ordered = re.match(r"^(\d+)[.、)]\s*(.*)", stripped)
        if ordered:
            _add_body_paragraph(doc, f"{ordered.group(1)}. {ordered.group(2)}", indent=first_line_indent)
            i += 1
            continue

        if re.match(r"^[-*]\s", stripped):
            _add_body_paragraph(doc, "• " + stripped[2:], indent=first_line_indent)
            i += 1
            continue

        if stripped in ("此致", "此致：", "此致:"):
            p = doc.add_paragraph("此致")
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue

        _add_body_paragraph(doc, stripped, indent=first_line_indent)
        i += 1


def build_document_from_markdown(
    content: str,
    *,
    title: str | None = None,
    author: str = "劳权智助",
) -> Document:
    doc = Document()
    _setup_page_simple(doc)
    _configure_document_styles(doc)

    core = doc.core_properties
    core.title = _sanitize_xml_text(title or "法律文书")
    core.author = author
    core.created = datetime.now()
    core.modified = datetime.now()

    body = (content or "").strip()
    if not body:
        body = "（文档内容为空）"

    if not body.lstrip().startswith("#") and title:
        h = doc.add_heading(_sanitize_xml_text(title), level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_markdown_to_document(doc, body, first_line_indent=True)

    _add_page_number(doc)

    footer = doc.add_paragraph("本文书由劳权智助辅助生成，提交前请核对事实与法律依据。")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def safe_docx_filename(doc_id: Any, title: str) -> str:
    safe_title = _SAFE_FILENAME_RE.sub("", title or "")
    if not safe_title:
        safe_title = f"document_{doc_id}"
    return f"{doc_id}_{safe_title}.docx"
