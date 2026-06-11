"""python-docx 直接填充引擎。

直接操作 .docx 模板文件中的占位符，用变量值替换。
不经过 Markdown 中间层，格式 100% 由模板控制。

流程：变量 JSON → 加载 .docx 模板 → 替换 {variable} → 输出 .docx
"""

from __future__ import annotations

import copy
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── 字体配置（与 create_templates.py 一致）──
_TITLE_FONT = "方正小标宋简体"
_HEADING1_FONT = "黑体"
_HEADING2_FONT = "楷体"
_BODY_FONT = "仿宋"
_EN_FONT = "Times New Roman"

_BODY_SIZE = Pt(16)
_LINE_SPACING = Pt(28.8)


def _set_east_asia_font(run, name: str) -> None:
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False) -> None:
    run.font.name = _EN_FONT
    run.font.size = size
    run.bold = bold
    _set_east_asia_font(run, font_name)


# ── 模板文件路径 ──
_TEMPLATE_DIR = Path(__file__).parent / "templates"

_TEMPLATE_FILES: dict[str, str] = {
    "application": "application_template.docx",
    "complaint": "complaint_template.docx",
    "labor_supervision": "labor_supervision_template.docx",
}


def get_template_path(doc_type: str) -> Path | None:
    """获取文书类型对应的模板文件路径。"""
    filename = _TEMPLATE_FILES.get(doc_type)
    if not filename:
        return None
    path = _TEMPLATE_DIR / filename
    return path if path.exists() else None


class DirectDocxFillEngine:
    """直接填充 .docx 模板中的占位符。"""

    def __init__(self):
        self._cache: dict[str, Document] = {}

    def _load_template(self, doc_type: str) -> Document | None:
        """加载模板（带缓存）。"""
        if doc_type in self._cache:
            return self._cache[doc_type]

        path = get_template_path(doc_type)
        if not path:
            return None

        doc = Document(str(path))
        self._cache[doc_type] = doc
        return doc

    def fill(
        self,
        doc_type: str,
        variables: dict[str, Any],
        *,
        evidence_list: list[dict[str, str]] | None = None,
    ) -> bytes:
        """填充模板并返回 DOCX 字节流。

        Args:
            doc_type: 文书类型（application / complaint / labor_supervision）
            variables: 变量字典 {key: value}
            evidence_list: 证据清单 [{name, purpose, pages}, ...]

        Returns:
            DOCX 文件的 bytes
        """
        doc = self._load_template(doc_type)
        if doc is None:
            raise ValueError(f"No template found for doc_type: {doc_type}")

        # 1. 替换段落中的占位符
        self._replace_in_paragraphs(doc, variables)

        # 2. 替换表格中的占位符
        self._replace_in_tables(doc, variables)

        # 3. 填充证据表格（仅 evidence_list 类型文书）
        # 证据是辅助认定事实的材料，不应出现在仲裁申请书/起诉状等文书中
        if evidence_list and doc_type == "evidence_list":
            self._fill_evidence_table(doc, evidence_list)

        # 4. 保存为 bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _replace_in_paragraphs(self, doc: Document, variables: dict[str, Any]) -> None:
        """遍历所有段落，替换 {variable} 占位符。"""
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, variables)

    def _replace_placeholders_in_paragraph(
        self, paragraph, variables: dict[str, Any]
    ) -> None:
        """替换段落中的占位符，保持原有格式。"""
        full_text = paragraph.text

        # 检查是否包含任何占位符
        has_placeholder = False
        for key in variables:
            if f"{{{key}}}" in full_text:
                has_placeholder = True
                break

        if not has_placeholder:
            return

        # 收集所有 runs 的文本
        runs = list(paragraph.runs)
        if not runs:
            return

        # 构建完整的文本映射
        # 策略：找到包含占位符的 run，直接替换
        for key, value in variables.items():
            placeholder = f"{{{key}}}"
            if placeholder not in full_text:
                continue

            str_value = str(value) if value is not None else ""

            # 找到包含占位符的 run
            replaced = False
            for run in runs:
                if placeholder in run.text:
                    # 空值：替换为"[待填写]"
                    if not str_value:
                        run.text = run.text.replace(placeholder, "[待填写]")
                    else:
                        # 检查值是否包含多段落分隔符
                        if "\n\n" in str_value:
                            # 多段落内容：在当前段落替换为第一段，后续段落插入到后面
                            paragraphs_parts = str_value.split("\n\n")
                            run.text = run.text.replace(placeholder, paragraphs_parts[0])
                            _set_run_font(run, _BODY_FONT, _BODY_SIZE, bold=run.bold)
                            
                            # 在当前段落之后插入后续段落
                            parent = paragraph._element.getparent()
                            insert_after = paragraph._element
                            for part in paragraphs_parts[1:]:
                                if part.strip():  # 跳过空段落
                                    new_p = OxmlElement("w:p")
                                    # 复制段落格式
                                    pPr = paragraph._element.find(qn("w:pPr"))
                                    if pPr is not None:
                                        new_pPr = copy.deepcopy(pPr)
                                        new_p.insert(0, new_pPr)
                                    # 添加 run
                                    new_r = OxmlElement("w:r")
                                    rPr = OxmlElement("w:rPr")
                                    # 设置字体
                                    rFonts = OxmlElement("w:rFonts")
                                    rFonts.set(qn("w:ascii"), _EN_FONT)
                                    rFonts.set(qn("w:hAnsi"), _EN_FONT)
                                    rFonts.set(qn("w:eastAsia"), _BODY_FONT)
                                    rPr.append(rFonts)
                                    sz = OxmlElement("w:sz")
                                    sz.set(qn("w:val"), str(int(_BODY_SIZE.pt * 2)))
                                    rPr.append(sz)
                                    szCs = OxmlElement("w:szCs")
                                    szCs.set(qn("w:val"), str(int(_BODY_SIZE.pt * 2)))
                                    rPr.append(szCs)
                                    new_r.append(rPr)
                                    new_t = OxmlElement("w:t")
                                    new_t.set(qn("xml:space"), "preserve")
                                    new_t.text = part.strip()
                                    new_r.append(new_t)
                                    new_p.append(new_r)
                                    # 插入到当前段落之后
                                    insert_after.addnext(new_p)
                                    insert_after = new_p
                        else:
                            run.text = run.text.replace(placeholder, str_value)
                    _set_run_font(run, _BODY_FONT, _BODY_SIZE, bold=run.bold)
                    replaced = True
                    break

            if not replaced:
                # 占位符跨多个 runs，需要合并处理
                self._replace_cross_run_placeholder(paragraph, placeholder, str_value)

    def _replace_cross_run_placeholder(
        self, paragraph, placeholder: str, value: str
    ) -> None:
        """处理占位符跨多个 runs 的情况。"""
        # 合并所有 run 文本
        full_text = paragraph.text
        if placeholder not in full_text:
            return

        # 清空所有 runs
        for run in paragraph.runs:
            run.text = ""

        # 在第一个 run 中设置完整文本
        if paragraph.runs:
            paragraph.runs[0].text = full_text.replace(placeholder, value)
            _set_run_font(paragraph.runs[0], _BODY_FONT, _BODY_SIZE)
        else:
            run = paragraph.add_run(value)
            _set_run_font(run, _BODY_FONT, _BODY_SIZE)

    def _replace_in_tables(self, doc: Document, variables: dict[str, Any]) -> None:
        """遍历所有表格，替换单元格中的占位符。"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, variables)

    def _fill_evidence_table(
        self, doc: Document, evidence_list: list[dict[str, str]]
    ) -> None:
        """填充证据清单表格。

        找到文档中第一个包含"序号"表头的表格，
        用证据列表数据填充（保留表头行）。
        """
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                first_row_text = " ".join(
                    cell.text.strip() for cell in table.rows[0].cells
                )
                if "序号" in first_row_text:
                    target_table = table
                    break

        if target_table is None:
            return

        # 保留表头，清除示例行
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[-1]._tr)

        # 添加证据行
        for i, item in enumerate(evidence_list, 1):
            row = target_table.add_row()
            name = item.get("name", "")
            purpose = item.get("purpose", "")
            pages = item.get("pages", "")

            # 清理内容中的 | 字符
            name = name.replace("|", "｜").replace("\n", " ")
            purpose = purpose.replace("|", "｜").replace("\n", " ")

            cells = row.cells
            if len(cells) >= 4:
                cells[0].text = str(i)
                cells[1].text = name
                cells[2].text = purpose
                cells[3].text = pages

                # 设置字体
                for cell in cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(24)
                        for run in p.runs:
                            _set_run_font(run, _BODY_FONT, Pt(14))

    def fill_from_markdown(
        self,
        doc_type: str,
        variables: dict[str, Any],
        *,
        evidence_text: str = "",
    ) -> bytes:
        """从 Markdown 格式的证据文本中提取证据列表并填充。

        这是便捷方法，解析简单的 Markdown 证据列表。
        """
        evidence_list = self._parse_evidence_markdown(evidence_text)
        return self.fill(doc_type, variables, evidence_list=evidence_list)

    def _parse_evidence_markdown(
        self, evidence_text: str
    ) -> list[dict[str, str]]:
        """解析 Markdown 格式的证据列表。

        支持格式：
        1. 证据名称—证明目的
        2. 证据名称：证明目的
        3. 证据名称 证明目的
        """
        if not evidence_text:
            return []

        items = []
        for line in evidence_text.strip().split("\n"):
            line = line.strip()
            if not line:
                continue

            # 移除编号前缀
            line = re.sub(r"^\d+[.、)]\s*", "", line)

            # 尝试分割名称和目的
            name = line
            purpose = ""

            for sep in ("—", "－", "：", ":", "，证明"):
                if sep in line:
                    name, _, purpose = line.partition(sep)
                    name = name.strip()
                    purpose = purpose.strip()
                    break

            if "证明" in line and not purpose:
                idx = line.find("证明")
                name = line[:idx].strip("，、 ")
                purpose = line[idx:].strip()

            if name:
                items.append({
                    "name": name,
                    "purpose": purpose or "证明相关案件事实",
                    "pages": "",
                })

        return items


# ── 便捷函数 ──

def fill_docx_from_variables(
    doc_type: str,
    variables: dict[str, Any],
    *,
    evidence_list: list[dict[str, str]] | None = None,
) -> bytes:
    """便捷函数：直接用变量填充模板。"""
    engine = DirectDocxFillEngine()
    return engine.fill(doc_type, variables, evidence_list=evidence_list)


def fill_docx_from_markdown(
    doc_type: str,
    variables: dict[str, Any],
    *,
    evidence_text: str = "",
) -> bytes:
    """便捷函数：从 Markdown 证据文本填充模板。"""
    engine = DirectDocxFillEngine()
    return engine.fill_from_markdown(doc_type, variables, evidence_text=evidence_text)
