"""程序化生成三个标准法律文书 .docx 模板文件。

基于官方标准格式：
1. 劳动仲裁申请书 — 依据《劳动争议调解仲裁法》第28条
2. 民事起诉状（劳动争议）— 依据最高人民法院2024年示范文本
3. 劳动保障监察投诉书 — 依据各地人社局标准格式

所有模板使用法院标准排版参数。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─ 字体配置（法院标准） ─
_TITLE_FONT = "方正小标宋简体"
_TITLE_FONT_ALT = "FZXiaoBiaoSong-B05S"
_HEADING1_FONT = "黑体"
_HEADING1_FONT_ALT = "SimHei"
_HEADING2_FONT = "楷体"
_HEADING2_FONT_ALT = "KaiTi"
_BODY_FONT = "仿宋"
_BODY_FONT_ALT = "FangSong"
_EN_FONT = "Times New Roman"

_TITLE_SIZE = Pt(22)
_HEADING1_SIZE = Pt(16)
_HEADING2_SIZE = Pt(16)
_HEADING3_SIZE = Pt(14)
_BODY_SIZE = Pt(16)
_LINE_SPACING = Pt(28.8)
_FIRST_LINE_INDENT = Cm(0.74)


def _set_east_asia_font(run, name: str) -> None:
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False) -> None:
    run.font.name = _EN_FONT
    run.font.size = size
    run.bold = bold
    _set_east_asia_font(run, font_name)


def _setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def _setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = _EN_FONT
    normal.font.size = _BODY_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(normal, _BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    run = p.add_run(text)
    _set_run_font(run, _TITLE_FONT, _TITLE_SIZE, bold=True)


def _add_section_heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.first_line_indent = None
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)

    if level == 2:
        font_name = _HEADING1_FONT
        size = _HEADING1_SIZE
    else:
        font_name = _HEADING2_FONT
        size = _HEADING2_SIZE

    run = p.add_run(text)
    _set_run_font(run, font_name, size, bold=True)


def _add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_after = Pt(4)
    if indent:
        pf.first_line_indent = _FIRST_LINE_INDENT
    else:
        pf.first_line_indent = None
    run = p.add_run(text)
    _set_run_font(run, _BODY_FONT, _BODY_SIZE)


def _add_no_indent_paragraph(doc: Document, text: str, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.first_line_indent = None
    pf.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, _BODY_FONT, _BODY_SIZE)


def _add_party_table(doc: Document, role: str, fields: list[tuple[str, str]]) -> None:
    """添加当事人信息表格（两列：标签 + 值）。"""
    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for ri, (label, placeholder) in enumerate(fields):
        # 标签列
        cell0 = table.cell(ri, 0)
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p0.paragraph_format.line_spacing = Pt(24)
        run0 = p0.add_run(f"{role}{label}：")
        _set_run_font(run0, _BODY_FONT, _BODY_SIZE, bold=True)

        # 值列
        cell1 = table.cell(ri, 1)
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.line_spacing = Pt(24)
        run1 = p1.add_run(placeholder)
        _set_run_font(run1, _BODY_FONT, _BODY_SIZE)

    # 添加空行
    doc.add_paragraph()


def _add_evidence_table(doc: Document) -> None:
    """添加证据清单表格（四列）。"""
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    headers = ["序号", "证据名称", "证明目的", "页数"]
    for ci, header in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run(header)
        _set_run_font(run, _BODY_FONT, Pt(14), bold=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9E2F3")
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)

    # 示例行
    for ci, text in enumerate(["1", "{evidence_name}", "{evidence_purpose}", ""]):
        cell = table.cell(1, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run(text)
        _set_run_font(run, _BODY_FONT, Pt(14))


def _add_closing(doc: Document, addressee: str, date: str = "{sign_date}") -> None:
    """添加落款（此致 + 机构 + 签名 + 日期）。"""
    _add_no_indent_paragraph(doc, "此致")
    _add_no_indent_paragraph(doc, addressee)
    doc.add_paragraph()
    _add_no_indent_paragraph(doc, "申请人：（签名）", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_no_indent_paragraph(doc, date, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ── 模板1：劳动仲裁申请书 ──

def create_application_template(output_dir: Path) -> Path:
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    _add_title(doc, "劳动人事争议仲裁申请书")

    _add_party_table(doc, "", [
        ("姓名", "{applicant_name}"),
        ("性别", "{applicant_gender}"),
        ("出生日期", "{applicant_birth}"),
        ("身份证号", "{applicant_id}"),
        ("住所", "{applicant_address}"),
        ("联系电话", "{applicant_phone}"),
        ("职业", "{applicant_job}"),
    ])

    _add_party_table(doc, "", [
        ("被申请人名称", "{respondent_name}"),
        ("统一社会信用代码", "{respondent_usci}"),
        ("住所", "{respondent_address}"),
        ("法定代表人", "{respondent_legal_rep}"),
        ("职务", "{respondent_position}"),
    ])

    _add_section_heading(doc, "一、仲裁请求")
    _add_body_paragraph(doc, "{requests}")

    _add_section_heading(doc, "二、事实与理由")
    _add_body_paragraph(doc, "{facts_and_reasons}")

    _add_section_heading(doc, "三、证据目录")
    _add_evidence_table(doc)

    _add_closing(doc, "{arbitration_commission}")

    filepath = output_dir / "application_template.docx"
    doc.save(str(filepath))
    return filepath


# ── 模板2：民事起诉状（劳动争议） ──

def create_complaint_template(output_dir: Path) -> Path:
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    _add_title(doc, "民事起诉状")

    _add_party_table(doc, "原告", [
        ("姓名", "{plaintiff_name}"),
        ("性别", "{plaintiff_gender}"),
        ("出生日期", "{plaintiff_birth}"),
        ("身份证号", "{plaintiff_id}"),
        ("住所", "{plaintiff_address}"),
        ("联系电话", "{plaintiff_phone}"),
    ])

    _add_party_table(doc, "被告", [
        ("名称", "{defendant_name}"),
        ("统一社会信用代码", "{defendant_usci}"),
        ("住所", "{defendant_address}"),
        ("法定代表人", "{defendant_legal_rep}"),
        ("职务", "{defendant_position}"),
    ])

    _add_section_heading(doc, "一、诉讼请求")
    _add_body_paragraph(doc, "{claims}")

    _add_section_heading(doc, "二、事实和理由")

    _add_section_heading(doc, "（一）劳动合同签订情况", level=3)
    _add_body_paragraph(doc, "{contract_info}")

    _add_section_heading(doc, "（二）劳动合同履行情况", level=3)
    _add_body_paragraph(doc, "{employment_facts}")

    _add_section_heading(doc, "（三）解除或终止劳动关系情况", level=3)
    _add_body_paragraph(doc, "{termination_facts}")

    _add_section_heading(doc, "（四）劳动仲裁情况", level=3)
    _add_body_paragraph(doc, "{arbitration_info}")

    _add_section_heading(doc, "三、证据清单")
    _add_evidence_table(doc)

    _add_no_indent_paragraph(doc, "此致")
    _add_no_indent_paragraph(doc, "{court_name}")
    doc.add_paragraph()
    _add_no_indent_paragraph(doc, "起诉人：（签名）", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_no_indent_paragraph(doc, "{sign_date}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    filepath = output_dir / "complaint_template.docx"
    doc.save(str(filepath))
    return filepath


# ── 模板3：劳动保障监察投诉书 ──

def create_labor_supervision_template(output_dir: Path) -> Path:
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    _add_title(doc, "劳动保障监察投诉书")

    _add_party_table(doc, "投诉人", [
        ("姓名", "{complainant_name}"),
        ("性别", "{complainant_gender}"),
        ("年龄", "{complainant_age}"),
        ("职业", "{complainant_job}"),
        ("工作单位", "{complainant_workplace}"),
        ("住所和联系方式", "{complainant_address}"),
    ])

    _add_party_table(doc, "被投诉人", [
        ("名称（用人单位）", "{employer_name}"),
        ("单位地址", "{employer_address}"),
        ("法定代表人或负责人姓名", "{employer_rep}"),
        ("联系电话", "{employer_phone}"),
    ])

    _add_section_heading(doc, "投诉人劳动保障合法权益受到侵害的事实")
    _add_body_paragraph(doc, "{facts}")

    _add_section_heading(doc, "请求事项")
    _add_body_paragraph(doc, "{items}")

    doc.add_paragraph()
    _add_no_indent_paragraph(doc, "投诉人签名：", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_no_indent_paragraph(doc, "{sign_date}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    filepath = output_dir / "labor_supervision_template.docx"
    doc.save(str(filepath))
    return filepath


# ── 批量生成 ──

def create_all_templates(output_dir: Path | None = None) -> dict[str, Path]:
    """生成所有三个标准模板文件。"""
    if output_dir is None:
        output_dir = Path(__file__).parent

    output_dir.mkdir(parents=True, exist_ok=True)

    results = {}
    results["application"] = create_application_template(output_dir)
    results["complaint"] = create_complaint_template(output_dir)
    results["labor_supervision"] = create_labor_supervision_template(output_dir)

    return results


if __name__ == "__main__":
    templates = create_all_templates()
    for name, path in templates.items():
        print(f"[OK] {name}: {path}")
