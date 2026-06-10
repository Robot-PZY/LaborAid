"""HTML → DOCX 转换引擎。

核心思路：前端预览和 Word 导出共用同一份 HTML，
确保预览和导出的格式 100% 一致。

流程：Markdown → HTML（court_markdown.py）→ DOCX（本模块）
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─ 字体配置（法院标准） ──
_TITLE_FONT = "方正小标宋简体"
_TITLE_FONT_ALT = "FZXiaoBiaoSong-B05S"
_HEADING1_FONT = "黑体"
_HEADING1_FONT_ALT = "SimHei"
_HEADING2_FONT = "楷体"
_HEADING2_FONT_ALT = "KaiTi"
_BODY_FONT = "仿宋"
_BODY_FONT_ALT = "FangSong"
_EN_FONT = "Times New Roman"

_TITLE_SIZE = Pt(22)
_HEADING1_SIZE = Pt(16)
_HEADING2_SIZE = Pt(16)
_HEADING3_SIZE = Pt(14)
_BODY_SIZE = Pt(16)
_LINE_SPACING = Pt(28.8)
_FIRST_LINE_INDENT = Cm(0.74)


def _set_east_asia_font(run, name: str) -> None:
    """设置 run 的东亚字体（不覆盖拉丁字体名）。"""
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False) -> None:
    """设置 run 的字体、字号、加粗。"""
    run.font.name = _EN_FONT
    run.font.size = size
    run.bold = bold
    _set_east_asia_font(run, font_name)


def _setup_page(doc: Document) -> None:
    """A4 纸张，法院标准页边距。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def _setup_styles(doc: Document) -> None:
    """统一 Normal 样式。"""
    normal = doc.styles["Normal"]
    normal.font.name = _EN_FONT
    normal.font.size = _BODY_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(normal, _BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_page_number(doc: Document) -> None:
    """在页脚添加居中页码（- X - 格式）。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        run_pre = p.add_run("- ")
        _set_run_font(run_pre, _BODY_FONT, Pt(10))

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_field = p.add_run()
        run_field._element.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run_instr = p.add_run()
        run_instr._element.append(instr_text)
        _set_run_font(run_instr, _BODY_FONT, Pt(10))

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end = p.add_run()
        run_end._element.append(fld_char_end)

        run_suf = p.add_run(" -")
        _set_run_font(run_suf, _BODY_FONT, Pt(10))


def _get_font_for_tag(tag_name: str, level: int = 0) -> tuple[str, Pt]:
    """根据 HTML 标签返回 (字体名, 字号)。"""
    if tag_name == "h1":
        return _TITLE_FONT, _TITLE_SIZE
    elif tag_name == "h2":
        return _HEADING1_FONT, _HEADING1_SIZE
    elif tag_name == "h3":
        return _HEADING2_FONT, _HEADING2_SIZE
    elif tag_name == "h4":
        return _BODY_FONT, _HEADING3_SIZE
    return _BODY_FONT, _BODY_SIZE


def _parse_inline_elements(p, element: Tag, font_name: str, size: Pt) -> None:
    """递归解析行内元素（strong, em, span 等），添加到段落中。"""
    for child in element.children:
        if isinstance(child, str):
            # 纯文本
            text = child
            if text:
                run = p.add_run(text)
                _set_run_font(run, font_name, size)
        elif isinstance(child, Tag):
            if child.name == "strong" or child.name == "b":
                run = p.add_run(child.get_text())
                _set_run_font(run, font_name, size, bold=True)
            elif child.name == "em" or child.name == "i":
                run = p.add_run(child.get_text())
                _set_run_font(run, font_name, size)
                run.italic = True
            elif child.name == "span":
                # 检查是否有特殊 class（如 template-hint）
                css_class = child.get("class", [])
                if "template-hint" in css_class:
                    run = p.add_run(child.get_text())
                    _set_run_font(run, font_name, size)
                    run.font.color.rgb = RGBColor(200, 0, 0)
                    run.bold = True
                else:
                    _parse_inline_elements(p, child, font_name, size)
            elif child.name == "br":
                # 换行
                run = p.add_run("\n")
                _set_run_font(run, font_name, size)
            else:
                # 其他标签，递归处理
                _parse_inline_elements(p, child, font_name, size)


def _add_paragraph_from_element(doc: Document, element: Tag) -> None:
    """从 HTML 元素创建 Word 段落。"""
    css_class = element.get("class", [])
    text = element.get_text().strip()

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING

    # 判断对齐方式
    if "center" in css_class or element.get("style", "").find("text-align:center") >= 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "right-align" in css_class:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.first_line_indent = None
    elif "no-indent" in css_class:
        pf.first_line_indent = None
    else:
        pf.first_line_indent = _FIRST_LINE_INDENT

    # 间距
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

    font_name, size = _get_font_for_tag("p")
    _parse_inline_elements(p, element, font_name, size)


def _add_heading_from_element(doc: Document, element: Tag) -> None:
    """从 HTML heading 元素创建 Word 标题。"""
    level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
    level = level_map.get(element.name, 2)

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _LINE_SPACING
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)

    font_name, size = _get_font_for_tag(element.name, level)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(12)
    else:
        pf.first_line_indent = None

    _parse_inline_elements(p, element, font_name, size)
    # 标题加粗
    for run in p.runs:
        run.bold = True


def _add_table_from_element(doc: Document, element: Tag) -> None:
    """从 HTML table 元素创建 Word 表格。"""
    rows = element.find_all("tr")
    if not rows:
        return

    # 解析所有行
    table_data: list[list[str]] = []
    for row in rows:
        cells = row.find_all(["th", "td"])
        row_data = [cell.get_text().strip() for cell in cells]
        if row_data:
            table_data.append(row_data)

    if not table_data:
        return

    max_cols = max(len(r) for r in table_data)
    if max_cols == 0:
        return

    # 规范化行数
    for i in range(len(table_data)):
        if len(table_data[i]) < max_cols:
            table_data[i] = table_data[i] + [""] * (max_cols - len(table_data[i]))

    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True

    for ri, row_data in enumerate(table_data):
        for ci in range(max_cols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(24)
            run = p.add_run(cell_text)
            _set_run_font(run, _BODY_FONT, Pt(14))
            if ri == 0:
                run.bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9E2F3")
                shading.set(qn("w:val"), "clear")
                cell._element.get_or_add_tcPr().append(shading)


def _add_list_from_element(doc: Document, element: Tag, ordered: bool = False) -> None:
    """从 HTML ol/ul 元素创建 Word 列表。"""
    items = element.find_all("li", recursive=False)
    for idx, item in enumerate(items):
        text = item.get_text().strip()
        if not text:
            continue

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = _LINE_SPACING
        pf.first_line_indent = _FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)

        font_name, size = _get_font_for_tag("p")

        if ordered:
            prefix = f"{idx + 1}. "
        else:
            prefix = "• "

        run = p.add_run(prefix)
        _set_run_font(run, font_name, size)

        _parse_inline_elements(p, item, font_name, size)


def html_to_docx_bytes(html_content: str, *, title: str | None = None) -> bytes:
    """将 HTML 内容转换为 DOCX 字节流。

    Args:
        html_content: 完整的 HTML 内容（含 court-document 容器）
        title: 文档标题（可选）

    Returns:
        DOCX 文件的 bytes
    """
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    # 设置文档属性
    core = doc.core_properties
    core.title = title or "法律文书"
    core.author = "劳权智助"

    # 解析 HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # 找到主容器
    container = soup.find("div", class_="court-document")
    if not container:
        container = soup

    # 遍历所有子元素
    for element in container.children:
        if not isinstance(element, Tag):
            continue

        tag = element.name

        if tag in ("h1", "h2", "h3", "h4"):
            _add_heading_from_element(doc, element)
        elif tag == "p":
            _add_paragraph_from_element(doc, element)
        elif tag == "table":
            _add_table_from_element(doc, element)
        elif tag == "ol":
            _add_list_from_element(doc, element, ordered=True)
        elif tag == "ul":
            _add_list_from_element(doc, element, ordered=False)
        elif tag == "br":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        elif tag == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 40)
            _set_run_font(run, _BODY_FONT, _BODY_SIZE)
            run.font.color.rgb = RGBColor(128, 128, 128)

    # 添加页码
    _add_page_number(doc)

    # 底部声明
    footer_p = doc.add_paragraph("本文书由劳权智助辅助生成，提交前请核对事实与法律依据。")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 保存为 bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
