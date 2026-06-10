from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_simsun_cell(cell):
    """设置单元格内所有文本的字体为 SimSun（宋体）。"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'SimSun'  # 设置字体为宋体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # 确保中文字体设置为宋体
            run.font.size = Pt(12)  # 可根据需要设置字体大小

# def set_paragraph_line_spacing_cell(cell, spacing_pt):
#     """设置单元格内所有段落的行间距为固定值。"""
#     for paragraph in cell.paragraphs:
#         p = paragraph._element
#         pPr = p.find(qn('w:pPr'))
#         if pPr is None:
#             pPr = OxmlElement('w:pPr')
#             p.append(pPr)
#         # 删除可能存在的原有行间距设置
#         existing_spacing = pPr.find(qn('w:spacing'))
#         if existing_spacing is not None:
#             pPr.remove(existing_spacing)
#         # 添加新的行间距设置
#         spacing = OxmlElement('w:spacing')
#         spacing.set(qn('w:line'), str(spacing_pt * 20))  # 行间距是磅数的20倍
#         spacing.set(qn('w:lineRule'), 'exact')  # 设置为固定行距
#         pPr.append(spacing)

def set_paragraph_line_spacing_cell(cell, spacing_pt):
    """设置单元格内所有段落的行间距。"""
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(spacing_pt)
        paragraph_format.line_spacing_rule = 1  # WD_LINE_SPACING.EXACTLY