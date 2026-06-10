from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.table import _Row  # 需要引入 _Row 类来重新封装 XML
from copy import deepcopy    # 引入深拷贝

def generate_ELdocx(evidence_name_list, evidence_source_list, evidence_object_list, all_names_sources_objects, template_docx):
    """
    【DeepCopy 修复版】
    解决逻辑：使用 deepcopy 复制模板行的 XML，确保每次生成的新行都是"干净"且未合并的。
    避免 add_row() 继承上一行已合并结构导致后续行索引越界的问题。
    """
    doc = Document(template_docx)
    table = doc.tables[0]
    
    # --- 1. 确定起始锚点 ---
    start_row_index = 2
    
    # 容错：防止表格行数不够
    while len(table.rows) <= start_row_index:
        table.add_row()
        
    # 获取模板行（始终以这一行为基准进行复制）
    # 注意：模板行本身应该是在 Word 模板里没有合并单元格的（或者符合你代码逻辑的）
    template_row = table.rows[start_row_index]
    
    # last_row_processed 是游标
    last_row_processed = template_row

    # --- 2. 遍历填数 ---
    for i, item_data in enumerate(all_names_sources_objects):
        e_name = item_data.get('evidence_name', '')
        e_source = item_data.get('evidence_source', '')
        e_object = item_data.get('evidence_object', '')

        current_row = None
        
        if i == 0:
            # === 第一条数据 ===
            # 直接使用现有的模板行
            current_row = last_row_processed
        else:
            # === 后续数据 ===
            # 1. 【核心修复】深拷贝模板行的 XML 元素
            # 这样保证新行和模板行一模一样（未合并状态），不受上一行影响
            new_tr = deepcopy(template_row._tr)
            
            # 2. 将新行直接插入到 last_row_processed 的后面
            # 使用 lxml 的 addnext 方法，比计算 index 更稳健
            last_row_processed._tr.addnext(new_tr)
            
            # 3. 将 XML 包装回 python-docx 的 Row 对象
            current_row = _Row(new_tr, table)

            # 4. 执行合并 (因为是深拷贝的干净行，cells[3] 必然存在，不会报错)
            current_row.cells[1].merge(current_row.cells[3])

        # --- 3. 填充内容 ---
        # 1. 序号
        set_cell_text(current_row.cells[0], str(i + 1))
        
        # 2. 证据名称
        set_cell_text(current_row.cells[1], str(e_name))
        
        # 3. 证据来源 (注意：合并后索引可能会变，但在 docx 中 cells 列表通常会保留原有长度的引用，除非完全重构)
        # 如果遇到索引错误，请尝试使用 current_row.cells[-2] 和 cells[-1]
        set_cell_text(current_row.cells[7], str(e_source))
        set_cell_text(current_row.cells[8], str(e_object))

        # --- 4. 更新游标 ---
        last_row_processed = current_row

    return doc

# --- 辅助函数 (保持不变) ---
def set_cell_text(cell, text):
    cell.text = '' 
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(12)
    run.font.bold = True
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER