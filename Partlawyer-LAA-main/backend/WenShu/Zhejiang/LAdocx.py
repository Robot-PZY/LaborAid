from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def set_run_style(run, font_name='仿宋', font_size=14, bold=False, underline=False, color=RGBColor(0, 0, 0)):
    """统一设置 Run 的样式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold
    run.underline = underline
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_styled_paragraph(doc, text="", space_after=5, line_spacing=1.5, bold=False, level=None):
    """辅助函数：添加带有特定格式的段落"""
    if level is not None:
        p = doc.add_heading(level=level)
    else:
        p = doc.add_paragraph()
    
    # 设置段落间距
    p.paragraph_format.space_after = Pt(space_after)
    
    # 设置行距 (标题通常不需要多倍行距，普通段落需要)
    if level is None:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_spacing

    if text:
        run = p.add_run(text)
        set_run_style(run, bold=bold, font_size=18 if level == 2 else 14)
    return p

def generate_LAdocx(claimant, gender, birthday_year, birthday_month, birthday_day, phone_claimant, id_number, nationality, occupation, registered_address, residence_address, respondent, legal_representative, phone_respondent, respondent_address, social_unity_credit_code, recourse_amount, debt_year_pre, debt_month_pre, debt_day_pre, debt_year_sub, debt_month_sub, debt_day_sub, hiredate_year, hiredate_month, hiredate_day, monthly_salary_standard, unpaid_year_pre, unpaid_month_pre, unpaid_day_pre, unpaid_year_sub, unpaid_month_sub, unpaid_day_sub, respondent_city, all_date_ranges):
    
    doc = Document()

    # --- 1. 设置全局默认字体 ---
    # 这样就不需要为每个非加粗/非下划线的普通文本重复设置字体了
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.color.rgb = RGBColor(0, 0, 0)

    # --- 2. 标题 ---
    p_title = add_styled_paragraph(doc, "追索劳动报酬仲裁申请书", space_after=30, level=2, bold=True)
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 3. 申请人信息 ---
    add_styled_paragraph(doc, f"申请人：{claimant}", space_after=22.5, level=3, bold=True)
    
    # 使用列表配置普通字段，减少代码量
    applicant_fields = [
        f"性别：{gender}",
        f"出生日期：{birthday_year}年 {birthday_month}月 {birthday_day}日",
        f"联系电话：{phone_claimant}",
        f"身份证号：{id_number}",
        f"民族：{nationality}",
        f"职业：{occupation}",
        f"户籍地：{registered_address}",
        f"居住地：{residence_address}"
    ]
    for field in applicant_fields:
        add_styled_paragraph(doc, field)

    # --- 4. 被申请人信息 ---
    add_styled_paragraph(doc, f"被申请人：{respondent}    (公司名称)", space_after=22.5, level=3, bold=True)
    
    respondent_fields = [
        f"法定代表人：{legal_representative}",
        f"联系电话：{phone_respondent}",
        f"地址：{respondent_address}"
    ]
    for field in respondent_fields:
        add_styled_paragraph(doc, field)
    
    # 统一社会信用代码单独处理间距
    add_styled_paragraph(doc, f"统一社会信用代码：{social_unity_credit_code}", space_after=22.5)

    # --- 5. 请求事项 ---
    add_styled_paragraph(doc, "请求事项：", space_after=22.5, level=3, bold=True)

    # 请求金额段落
    p_req = add_styled_paragraph(doc, space_after=22.5) # 先建空段落
    # 拼接文本：普通文本 -> 下划线变量 -> 普通文本
    segments = [
        ("裁令被申请人向申请人支付劳动报酬", False),
        (f"   {recourse_amount}   ", True),
        ("元", False)
    ]
    for text, is_underline in segments:
        run = p_req.add_run(text)
        set_run_style(run, underline=is_underline)

    # 追索起止时间说明
    p_calc = add_styled_paragraph(doc, space_after=50)
    
    # 定义基础时间段列表
    date_segments = [
        ("（工资标准计算：日薪n元×欠薪天数，", False),
        (f" {debt_year_pre} ", True), ("年", False),
        (f" {debt_month_pre} ", True), ("月", False),
        (f" {debt_day_pre} ", True), ("日", False),
        ("计算至", False),
        (f" {debt_year_sub} ", True), ("年", False),
        (f" {debt_month_sub} ", True), ("月", False),
        (f" {debt_day_sub} ", True), ("日", False)
    ]

    # 处理额外的日期范围 (如果有)
    if all_date_ranges:
        for dr in all_date_ranges:
            date_segments.append(("、", False))
            date_segments.extend([
                (f" {dr['start_year']} ", True), ("年", False),
                (f" {dr['start_month']} ", True), ("月", False),
                (f" {dr['start_day']} ", True), ("日", False),
                ("计算至", False),
                (f" {dr['end_year']} ", True), ("年", False),
                (f" {dr['end_month']} ", True), ("月", False),
                (f" {dr['end_day']} ", True), ("日", False)
            ])
    
    date_segments.append(("）", False))

    # 渲染时间段
    for text, is_underline in date_segments:
        run = p_calc.add_run(text)
        set_run_style(run, underline=is_underline)

    # --- 6. 事实和理由 ---
    add_styled_paragraph(doc, "事实和理由：", space_after=22.5, level=3, bold=True)

    # 构造长段落
    p_facts = add_styled_paragraph(doc, space_after=12, line_spacing=2)
    
    facts_content = [
        ("申请人自（入职时间）", False),
        (f"  {hiredate_year}  ", True), ("年", False),
        (f"  {hiredate_month}  ", True), ("月", False),
        (f"  {hiredate_day}  ", True), ("日起，与被申请人建立劳动关系，工作岗位为", False),
        (f"   {monthly_salary_standard}元   ", True), ("月工资标准。被申请人在（少发、未发劳动报酬起止时间）", False),
        # 初始欠薪时间
        (f"  {unpaid_year_pre}  ", True), ("年", False),
        (f"  {unpaid_month_pre}  ", True), ("月", False),
        (f"  {unpaid_day_pre}  ", True), ("日至", False),
        (f"  {unpaid_year_sub}  ", True), ("年", False),
        (f"  {unpaid_month_sub}  ", True), ("月", False),
        (f"  {unpaid_day_sub}  ", True), ("日", False),
    ]

    # 处理额外的欠薪日期范围
    if all_date_ranges:
        for dr in all_date_ranges:
            facts_content.append(("、", False))
            facts_content.extend([
                (f"  {dr['start_year']}  ", True), ("年", False),
                (f"  {dr['start_month']}  ", True), ("月", False),
                (f"  {dr['start_day']}  ", True), ("日至", False),
                (f"  {dr['end_year']}  ", True), ("年", False),
                (f"  {dr['end_month']}  ", True), ("月", False),
                (f"  {dr['end_day']}  ", True), ("日", False),
            ])

    facts_content.append(("未足", False))
    facts_content.append(("额支付申请人劳动报酬，严重损害了申请人的合法利益。", False))

    # 渲染事实段落
    for text, is_underline in facts_content:
        run = p_facts.add_run(text)
        set_run_style(run, underline=is_underline)

    # --- 7. 结语 ---
    add_styled_paragraph(doc, "依据相关法律法规，特向贵委提出仲裁申请，望裁如所请。", space_after=27)
    add_styled_paragraph(doc, "此致", space_after=5)

    # 仲裁委名称
    p_committee = add_styled_paragraph(doc)
    run_city = p_committee.add_run(f"  {respondent_city}  ")
    set_run_style(run_city, underline=True)
    run_suffix = p_committee.add_run("市劳动人事争议仲裁委员会")
    set_run_style(run_suffix)

    # 底部提示语 (灰色字体)
    p_note1 = add_styled_paragraph(doc, "注意：根据《劳动仲裁法》第二十一条劳动争议由劳动合同履行地或者用人单位所在地的劳动争议仲裁委员会管辖", bold=True)
    p_note1.paragraph_format.space_before = Pt(10)
    p_note1.paragraph_format.space_after = Pt(10)
    p_note1.runs[0].font.color.rgb = RGBColor(108, 108, 108)

    # --- 8. 签字区域 ---
    p_sign = add_styled_paragraph(doc, "申请人（本人签字并捺手印）：")
    p_sign.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    p_date = add_styled_paragraph(doc, "年  月  日")
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 底部附录 (灰色字体)
    p_note2 = add_styled_paragraph(doc, "附（与起诉状同时提供）：证据、申请人身份证复印件、被申请人工商营业执照或工商档案主体信息", bold=True)
    p_note2.paragraph_format.space_before = Pt(11.25)
    p_note2.paragraph_format.space_after = Pt(11.25)
    p_note2.runs[0].font.color.rgb = RGBColor(108, 108, 108)

    return doc