from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# --- 样式辅助函数保持不变 ---

# 设置宋体三号加粗的函数
def set_font_simsun_bold(run):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)  # 三号字体
    run.font.bold = True    # 加粗

# 设置段落居中的函数
def set_paragraph_center(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置单元格垂直居中的函数
def set_cell_vertical_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# --- 核心生成函数修改 ---

def generate_ELdocx(evidence_names_list, evidence_contents_list, all_evidences_testifys, template_docx):
    """
    参数说明:
    evidence_names_list: 名字列表 (router传入的第1个参数) - 在这里没用到，因为我们用字典
    evidence_contents_list: 内容列表 (router传入的第2个参数) - 也没用到
    all_evidences_testifys: 字典列表 (router传入的第3个参数) - 主要用这个！
    template_docx: 模板路径
    """
    doc = Document(template_docx)
    table = doc.tables[0] # 获取第一个表格
    
    # 获取模板行的属性（用于复制高度）
    # 假设表格第1行是标题，第2行（索引1）是空的模板行
    template_row_index = 1
    template_row = table.rows[template_row_index]
    row_height = template_row.height # 记住高度

    # 遍历数据列表
    for i, item_data in enumerate(all_evidences_testifys):
        # 1. 决定是使用现有行，还是新增行
        if i == 0:
            # 如果是第1条数据，直接使用模板里预留的那一行
            row = table.rows[template_row_index]
        else:
            # 如果是第2条及以后，新增一行
            row = table.add_row()
            row.height = row_height # 复制高度
        
        # 2. 获取数据 (router 里的字典键名是 'name' 和 'content'，这里做了兼容)
        # 优先取 'name'，如果没有则取 'evidence_name'
        name_text = item_data.get('name') or item_data.get('evidence_name', '')
        content_text = item_data.get('content') or item_data.get('testify_content', '')
        
        # 3. 填充单元格 (序号、名称、内容)
        
        # --- 第一列：序号 ---
        cell_0 = row.cells[0]
        cell_0.text = '' # 清空原有内容（如果有）
        paragraph_0 = cell_0.paragraphs[0]
        run_0 = paragraph_0.add_run(str(i + 1)) # 填入序号 1, 2, 3...
        
        set_font_simsun_bold(run_0)        # 应用样式
        set_paragraph_center(paragraph_0)
        set_cell_vertical_center(cell_0)

        # --- 第二列：证据名称 ---
        cell_1 = row.cells[1]
        cell_1.text = ''
        paragraph_1 = cell_1.paragraphs[0]
        run_1 = paragraph_1.add_run(str(name_text)) # 确保是字符串
        
        set_font_simsun_bold(run_1)
        set_paragraph_center(paragraph_1)
        set_cell_vertical_center(cell_1)

        # --- 第三列：证明内容 ---
        cell_2 = row.cells[2]
        cell_2.text = ''
        paragraph_2 = cell_2.paragraphs[0]
        run_2 = paragraph_2.add_run(str(content_text))
        
        set_font_simsun_bold(run_2)
        set_paragraph_center(paragraph_2)
        set_cell_vertical_center(cell_2)

    return doc