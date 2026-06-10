from docx import Document
from utl.set_song_or_TNR import set_font_simsun
from utl.set_song_or_TNR import set_font_times_new_roman    
from utl.set_song_or_TNR import apply_font_based_on_content
from utl.paragraph_format import set_paragraph_center
from utl.paragraph_format import set_cell_vertical_center
from utl.paragraph_format import set_paragraph_line_spacing
from utl.split_text import split_text_by_length_and_newline
from utl.cell_format import set_font_simsun_cell

def generate_LAdocx(submitCity, applicantName, gender, birthYear, birthMonth, birthDay, idNumber, phone, address, addressOption, respondentName, respondentAddress, respondentAddressOption, legalRepName, legalRepPosition, legalRepContact, arbitrationRequest, requestCalculation, hireYear, hireMonth, hireDay, jobPosition, contractOption, contractStartYear, contractStartMonth, contractStartDay, contractEndYear, contractEndMonth, contractEndDay, workLocation, workTimeOption, workDaysPerWeek, workHoursPerDay, attendanceOption, attendanceMethod, salaryOption1, salaryOption2, salaryOption3, salaryOption4, initialSalary, salaryAdjustment, currentEmploymentStatus, leaveYear, leaveMonth, leaveDay, leaveReason, avgMonthlySalary, additionalFact, template_docx, addressAdditionalInput, respondentAddressAdditionalInput, workTimeAdditionalInput):
    
    # 加载 Word 文档
    doc = Document(template_docx)

    # -----------------------------------------------------------------------------
    # 定义内部辅助函数，优化代码复用
    # -----------------------------------------------------------------------------
    
    # 1. 日期填充函数 (年/月/日)
    def fill_date_cell(cell, year, month, day):
        cell.text = ""
        set_cell_vertical_center(cell)
        # 确保有段落
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        set_paragraph_center(p)
        
        # 年
        run = p.add_run(str(year))
        set_font_times_new_roman(run)
        run = p.add_run("年")
        set_font_simsun(run)
        
        # 月
        run = p.add_run("  " + str(month))
        set_font_times_new_roman(run)
        run = p.add_run("月")
        set_font_simsun(run)
        
        # 日
        run = p.add_run("  " + str(day))
        set_font_times_new_roman(run)
        run = p.add_run("日")
        set_font_simsun(run)

    # 2. 薪资选项处理函数 (修复了 UnboundLocalError)
    def process_salary_option(table, row, column, option_value, option_text):
        cell = table.cell(row, column)
        # 【关键修复】: 先获取文本，确保变量作用域覆盖整个函数
        cell_text = cell.text
        
        prefix = ""
        if option_value == 'yes':
            prefix = "☑"
        elif option_value == 'no':
            prefix = "□"
        else:
            return # 无操作

        # 执行替换
        if option_text in cell_text:
            # 防止重复添加 (简单校验)
            if ("☑" + option_text) not in cell_text and ("□" + option_text) not in cell_text:
                new_text = cell_text.replace(option_text, prefix + option_text)
                cell.text = new_text
                # 重新应用格式
                set_font_simsun_cell(cell)
                if cell.paragraphs:
                    set_paragraph_line_spacing(cell.paragraphs[0], 16)

    # -----------------------------------------------------------------------------
    # 主逻辑开始
    # -----------------------------------------------------------------------------
    
    # 查找并更新“致：”段落
    for paragraph in doc.paragraphs:
        if '致：劳动人事争议仲裁委员会' in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run('致：')
            set_font_simsun(run)
            
            city_run = paragraph.add_run("    " + submitCity + "    ")
            city_run.font.underline = True 
            set_font_simsun(city_run) 
            
            remaining_run = paragraph.add_run('劳动人事争议仲裁委员会')
            set_font_simsun(remaining_run) 
    
    
    # ***表一 (基本信息)***
    table = doc.tables[0]
    
    # --- 申请人信息 ---
    
    # 姓名
    cell_applicant_name = table.cell(0, 2)
    cell_applicant_name.text = ""
    name_run = cell_applicant_name.paragraphs[0].add_run(applicantName)
    set_font_simsun(name_run)
    set_paragraph_center(cell_applicant_name.paragraphs[0])
    set_cell_vertical_center(cell_applicant_name)
    
    # 性别
    cell_gender = table.cell(0, 4)
    cell_gender.text = ""
    gender_run = cell_gender.paragraphs[0].add_run(gender)
    set_font_simsun(gender_run)
    set_paragraph_center(cell_gender.paragraphs[0])
    set_cell_vertical_center(cell_gender)
    
    # 出生日期 (调用辅助函数)
    if birthYear and birthMonth and birthDay:
        cell_brith = table.cell(0, 7)
        fill_date_cell(cell_brith, birthYear, birthMonth, birthDay)
    
    # 身份证号
    cell_idNumber = table.cell(1, 4) 
    cell_idNumber.text = ""
    idNumber_run = cell_idNumber.paragraphs[0].add_run(idNumber)
    set_font_times_new_roman(idNumber_run)
    set_paragraph_center(cell_idNumber.paragraphs[0])
    set_cell_vertical_center(cell_idNumber)
    
    # 联系电话
    cell_phone = table.cell(1, 7)
    cell_phone.text = ""
    phone_run = cell_phone.paragraphs[0].add_run(phone)
    set_font_times_new_roman(phone_run)
    set_paragraph_center(cell_phone.paragraphs[0])
    set_cell_vertical_center(cell_phone)
    
    # 住所
    cell_address = table.cell(2, 2)
    cell_address.text=""
    paragraphs_num = 0
    apply_font_based_on_content(cell_address, address, paragraphs_num)
    set_paragraph_center(cell_address.paragraphs[0])
    set_cell_vertical_center(cell_address)
    
    # 申请人通讯方式选项
    cell_address_option = table.cell(3, 2)
    if addressOption == 'confirmation_of_partie':
        cell_address_option.text = "" 
        address_option_run = cell_address_option.paragraphs[0].add_run("☑与被申请人住所相同")
        set_font_simsun(address_option_run)
        
        paragraph_2 = cell_address_option.add_paragraph() 
        other_option_run = paragraph_2.add_run("□其他：")
        set_font_simsun(other_option_run)
        
    elif addressOption == 'other':
        cell_address_option.text = "" 
        address_option_run = cell_address_option.paragraphs[0].add_run("□与被申请人住所相同")
        set_font_simsun(address_option_run)
        
        paragraph_2 = cell_address_option.add_paragraph() 
        text = "☑其他：" + addressAdditionalInput
        apply_font_based_on_content(cell_address_option, text, 1) # 1 表示第二段
        set_paragraph_line_spacing(paragraph_2, 16)       
        
    # --- 被申请人信息 ---
    
    # 被申请人名称
    cell_respondent_name = table.cell(4, 2)
    cell_respondent_name.text = ""
    respondent_name_run = cell_respondent_name.paragraphs[0].add_run(respondentName)
    set_font_simsun(respondent_name_run)
    set_paragraph_center(cell_respondent_name.paragraphs[0])
    set_cell_vertical_center(cell_respondent_name)
    
    # 住所
    cell_respondent_address = table.cell(5, 2)
    cell_respondent_address.text = ""
    apply_font_based_on_content(cell_respondent_address, respondentAddress, 0)
    set_paragraph_center(cell_respondent_address.paragraphs[0])
    set_cell_vertical_center(cell_respondent_address)
       
    # 被申请人通讯方式选项
    cell_respondent_address_option = table.cell(6, 2)
    if respondentAddressOption == 'same_as_residence':
        cell_respondent_address_option.text = "" 
        run = cell_respondent_address_option.paragraphs[0].add_run("☑与被申请人住所相同")
        set_font_simsun(run)
        
        paragraph_2 = cell_respondent_address_option.add_paragraph() 
        run = paragraph_2.add_run("□其他：")
        set_font_simsun(run)
        
    elif respondentAddressOption == 'other':
        cell_respondent_address_option.text = ""
        run = cell_respondent_address_option.paragraphs[0].add_run("□与被申请人住所相同")
        set_font_simsun(run)
        
        paragraph_2 = cell_respondent_address_option.add_paragraph()
        text = "☑其他：" + respondentAddressAdditionalInput
        apply_font_based_on_content(cell_respondent_address_option, text, 1)
        set_paragraph_line_spacing(paragraph_2, 16)
    
    # 法定代表人姓名
    cell_legal_rep_name = table.cell(7, 2)
    cell_legal_rep_name.text = ""
    legal_rep_name_run = cell_legal_rep_name.paragraphs[0].add_run(legalRepName)
    set_font_simsun(legal_rep_name_run)
    set_paragraph_center(cell_legal_rep_name.paragraphs[0])
    set_cell_vertical_center(cell_legal_rep_name)
    
    # 法定代表人职务
    cell_legal_rep_position = table.cell(8, 2)
    cell_legal_rep_position.text = ""
    legal_rep_position_run = cell_legal_rep_position.paragraphs[0].add_run(legalRepPosition)
    set_font_simsun(legal_rep_position_run)
    set_paragraph_center(cell_legal_rep_position.paragraphs[0])
    set_cell_vertical_center(cell_legal_rep_position)
    
    # 法定代表人联系电话
    cell_legal_rep_contact = table.cell(8, 6)
    cell_legal_rep_contact.text = ""
    legal_rep_contact_run = cell_legal_rep_contact.paragraphs[0].add_run(legalRepContact)
    set_font_simsun(legal_rep_contact_run)
    set_paragraph_center(cell_legal_rep_contact.paragraphs[0])
    set_cell_vertical_center(cell_legal_rep_contact)


    # ***表二 (仲裁请求)***
    table2 = doc.tables[1]
    
    # 仲裁请求
    arbitration_request_content_row = 1
    arbitration_request_serial_number_row = 1
    
    for arbitrationRequest_item in arbitrationRequest:
        # 序号
        cell_request_idx = table2.cell(arbitration_request_content_row, 1)
        cell_request_idx.text = ""
        if arbitrationRequest_item:
            run = cell_request_idx.paragraphs[0].add_run(str(arbitration_request_serial_number_row) + '、')
            set_font_times_new_roman(run)
            set_paragraph_center(cell_request_idx.paragraphs[0])
        
        # 内容分割
        chunks = split_text_by_length_and_newline(arbitrationRequest_item, 28)
        for chunk in chunks:
            cell_content = table2.cell(arbitration_request_content_row, 2)
            cell_content.text = ""
            apply_font_based_on_content(cell_content, chunk, 0)
            set_paragraph_center(cell_content.paragraphs[0])
            set_cell_vertical_center(cell_content)
            
            arbitration_request_content_row += 1
        
        arbitration_request_serial_number_row += 1
    
    # 计算公式
    requestCalculation_row = 16
    for requestCalculation_item in requestCalculation:
        chunks = split_text_by_length_and_newline(requestCalculation_item, 30)
        for chunk in chunks:
            cell_calc = table2.cell(requestCalculation_row, 1)
            cell_calc.text = ""
            apply_font_based_on_content(cell_calc, chunk, 0)
            set_cell_vertical_center(cell_calc)
            requestCalculation_row += 1


    # ***表三 (劳动关系信息)***
    table3 = doc.tables[2]
    
    # 入职时间 (调用辅助函数)
    if hireYear and hireMonth and hireDay:
        cell_hire = table3.cell(1, 4)
        fill_date_cell(cell_hire, hireYear, hireMonth, hireDay)

    # 工作岗位
    cell_job_position = table3.cell(1, 6)
    cell_job_position.text = ""
    job_position_run = cell_job_position.paragraphs[0].add_run(jobPosition)
    set_font_simsun(job_position_run)
    set_paragraph_center(cell_job_position.paragraphs[0])
    set_cell_vertical_center(cell_job_position)
    
    # 合同选项
    cell_contract_option = table3.cell(1, 11)
    if contractOption == 'yes':
        cell_contract_option.text = "" 
        p1 = cell_contract_option.paragraphs[0]
        run = p1.add_run("☑有")
        set_font_simsun(run)
        set_paragraph_line_spacing(p1, 16)
    
        p2 = cell_contract_option.add_paragraph()
        run = p2.add_run("□无")
        set_font_simsun(run)
        set_paragraph_line_spacing(p2, 16)
        
    elif contractOption == 'no':
        cell_contract_option.text = "" 
        p1 = cell_contract_option.paragraphs[0]
        run = p1.add_run("□有")
        set_font_simsun(run)
        
        p2 = cell_contract_option.add_paragraph()
        run = p2.add_run("☑无")
        set_font_simsun(run)
        set_paragraph_line_spacing(p2, 16)

    # 合同日期
    if contractOption and contractStartYear and contractStartMonth and contractStartDay and contractEndYear and contractEndMonth and contractEndDay:
        cell_contract = table3.cell(2, 2)
        cell_contract.text = ""
        text = f"{contractStartYear}年{'  '}{contractStartMonth}月{'  '}{contractStartDay}日至{'  '}{contractEndYear}年{'  '}{contractEndMonth}月{'  '}{contractEndDay}日"
        apply_font_based_on_content(cell_contract, text, 0)
        set_paragraph_center(cell_contract.paragraphs[0])
        set_cell_vertical_center(cell_contract)

    # 工作地点
    cell_work_location = table3.cell(3, 2)
    cell_work_location.text = ""
    apply_font_based_on_content(cell_work_location, workLocation, 0)
    set_paragraph_center(cell_work_location.paragraphs[0])
    set_cell_vertical_center(cell_work_location)
    
    # 工作时间选项
    cell_work_time_option = table3.cell(4, 2)
    set_paragraph_center(cell_work_time_option.paragraphs[0])
    set_cell_vertical_center(cell_work_time_option)
    
    if workTimeOption == 'specific_time':
        cell_work_time_option.text = "" 
        
        run = cell_work_time_option.paragraphs[0].add_run("☑每周工作")
        set_font_simsun(run)
        
        # 处理数字字体
        run = cell_work_time_option.paragraphs[0].add_run('  ' + workDaysPerWeek + '  ')
        if workDaysPerWeek.isdigit():
            set_font_times_new_roman(run)
        else:
            set_font_simsun(run)
        run.underline = True
        
        run = cell_work_time_option.paragraphs[0].add_run("天，每天工作")
        set_font_simsun(run)
        
        run = cell_work_time_option.paragraphs[0].add_run('  ' + workHoursPerDay + '  ')
        if workHoursPerDay.isdigit():
            set_font_times_new_roman(run)
        else:
            set_font_simsun(run)
        run.underline = True
        
        run = cell_work_time_option.paragraphs[0].add_run("小时")
        set_font_simsun(run)
        
        # 新增一行 其他
        p2 = cell_work_time_option.add_paragraph()
        run = p2.add_run("□其他：")
        set_font_simsun(run)
      
    elif workTimeOption == 'other':
        cell_work_time_option.text = "" 
        run = cell_work_time_option.paragraphs[0].add_run("□每周工作_____天，每天工作_____小时")
        set_font_simsun(run)
        
        p2 = cell_work_time_option.add_paragraph()
        text = "☑其他：" + workTimeAdditionalInput
        apply_font_based_on_content(cell_work_time_option, text, 1)
    
    # 考勤选项
    cell_attendance_option = table3.cell(5, 2)
    if attendanceOption == 'yes':
        cell_attendance_option.text = ""
        p1 = cell_attendance_option.paragraphs[0]
        run = p1.add_run("☑是")
        set_font_simsun(run) 
        set_paragraph_line_spacing(p1, 16)
    
        p2 = cell_attendance_option.add_paragraph()
        run = p2.add_run("□否")
        set_font_simsun(run) 
        set_paragraph_line_spacing(p2, 16)
        
    elif attendanceOption == 'no':
        cell_attendance_option.text = "" 
        p1 = cell_attendance_option.paragraphs[0]
        run = p1.add_run("□是")
        set_font_simsun(run)
        
        p2 = cell_attendance_option.add_paragraph()
        run = p2.add_run("☑否")
        set_font_simsun(run)
        set_paragraph_line_spacing(p2, 16)

    # 考勤方式
    cell_attendance_method = table3.cell(5, 4)
    cell_attendance_method.text = ""
    run = cell_attendance_method.paragraphs[0].add_run(attendanceMethod)
    set_font_simsun(run)
    set_paragraph_center(cell_attendance_method.paragraphs[0])
    set_cell_vertical_center(cell_attendance_method)

    # 薪资选项 (使用修复后的函数)
    # 假设 Table3 的结构固定：第6行(index 5)，第8列(index 7)是支付方式，第11列(index 10)是签收方式
    process_salary_option(table3, 5, 7, salaryOption1, "现金")
    process_salary_option(table3, 5, 7, salaryOption2, "转账")
    process_salary_option(table3, 5, 10, salaryOption3, "需要签收")
    process_salary_option(table3, 5, 10, salaryOption4, "不需签收")

    # 初始薪资
    cell_initial_salary = table3.cell(6, 2)
    cell_initial_salary.text = ""
    apply_font_based_on_content(cell_initial_salary, initialSalary, 0)
    set_paragraph_center(cell_initial_salary.paragraphs[0])
    set_cell_vertical_center(cell_initial_salary)

    # 薪资调整
    cell_salary_adjustment = table3.cell(6, 6)
    cell_salary_adjustment.text = ""
    apply_font_based_on_content(cell_salary_adjustment, salaryAdjustment, 0)
    set_paragraph_center(cell_salary_adjustment.paragraphs[0])
    set_cell_vertical_center(cell_salary_adjustment)

    # 现是否在职
    cell_current_status = table3.cell(7, 2)
    # 获取原始文本用于替换
    cell_text = cell_current_status.text
    if currentEmploymentStatus == 'yes':
        if '是' in cell_text:
            new_text = cell_text.replace('□是', '☑是')
            cell_current_status.text = new_text
            set_font_simsun_cell(cell_current_status)
            if cell_current_status.paragraphs:
                set_paragraph_line_spacing(cell_current_status.paragraphs[0], 16)
    elif currentEmploymentStatus == 'no':
        new_text = cell_text.replace('□否', '☑否')       
        cell_current_status.text = new_text
        set_font_simsun_cell(cell_current_status)
        if cell_current_status.paragraphs:
            set_paragraph_line_spacing(cell_current_status.paragraphs[0], 16)
    
    set_paragraph_center(cell_current_status.paragraphs[0])
    set_cell_vertical_center(cell_current_status)

    # 离职日期 (调用辅助函数)
    if leaveYear and leaveMonth and leaveDay:
        cell_leave_date = table3.cell(7, 4)
        fill_date_cell(cell_leave_date, leaveYear, leaveMonth, leaveDay)

    # 离职原因
    cell_leave_reason = table3.cell(8, 9)
    cell_leave_reason.text = ""
    apply_font_based_on_content(cell_leave_reason, leaveReason, 0)
    set_paragraph_center(cell_leave_reason.paragraphs[0])
    set_cell_vertical_center(cell_leave_reason)

    # 平均月薪
    cell_avg_salary = table3.cell(9, 9)
    cell_avg_salary.text = ""
    fix_text = avgMonthlySalary + "月/元"
    apply_font_based_on_content(cell_avg_salary, fix_text, 0)
    set_paragraph_center(cell_avg_salary.paragraphs[0])
    set_cell_vertical_center(cell_avg_salary)


    # ***表四 (其他理由)***
    table4 = doc.tables[3]
    
    additionalFact_row = 1
    for additionalFact_item in additionalFact:
        chunks = split_text_by_length_and_newline(additionalFact_item, 30)
        for chunk in chunks:
            cell_fact = table4.cell(additionalFact_row, 1)
            cell_fact.text = ""
            apply_font_based_on_content(cell_fact, chunk, 0)
            set_cell_vertical_center(cell_fact)
            additionalFact_row += 1

    return doc