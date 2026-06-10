from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_UNDERLINE
from utl.num_to_upper import num_to_upper_main

def set_font_color_space(run, p, size):
    run.font.size = Pt(size)
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 设置段后距离为1.5磅
    p.paragraph_format.space_after = Pt(10.25)
    # 设置行距为最小值26.25磅
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(26.25)

def amount_add_underline(paragraph_name, fee_name, amount):
    run = paragraph_name.add_run(fee_name)
    set_font_color_space(run, paragraph_name, 14)
    run = paragraph_name.add_run("   " + (amount if amount else "") + "   ")
    run.underline = True
    set_font_color_space(run, paragraph_name, 14)
    run = paragraph_name.add_run("元")
    set_font_color_space(run, paragraph_name, 14)

def others_add_underline(paragraph_name, text, data):
    for word in text:
        if word in data and word not in ['selection_government_region', 'selection_application_region']:
            run = paragraph_name.add_run("   " + (data.get(word, "")) + "   ")
            run.underline = WD_UNDERLINE.SINGLE
            set_font_color_space(run, paragraph_name, 14)
        elif word in ['selection_government_region', 'selection_application_region']:
            run = paragraph_name.add_run(data.get(word, ""))
            set_font_color_space(run, paragraph_name, 14)
        else:
            run = paragraph_name.add_run(word)
            set_font_color_space(run, paragraph_name, 14)

def generate_WIAdocx(data):
    doc = Document()

    # 工伤仲裁申请书
    title = doc.add_paragraph()
    # 设置居中
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("工伤仲裁申请书")
    run.bold = True
    set_font_color_space(run, title, 18)

    # 空行
    doc.add_paragraph()

    # 申请人
    p_claimant = doc.add_paragraph()
    run = p_claimant.add_run("申请人：")
    run.bold = True
    set_font_color_space(run, p_claimant, 14)

    # 姓名
    pc_name = doc.add_paragraph()
    run = pc_name.add_run("姓名：" + data.get('claimant_name', ''))
    set_font_color_space(run, pc_name, 14)

    # 民族
    pc_nationality = doc.add_paragraph()
    run = pc_nationality.add_run("民族：" + data.get('claimant_nationality', ''))
    set_font_color_space(run, pc_nationality, 14)

    # 性别
    pc_gender = doc.add_paragraph()
    run = pc_gender.add_run("性别：" + data.get('gender', ''))
    set_font_color_space(run, pc_gender, 14)

    # 职业
    pc_occupation = doc.add_paragraph()
    run = pc_occupation.add_run("职业：" + data.get('occupation', ''))
    set_font_color_space(run, pc_occupation, 14)

    # 出生日期
    pc_birthday = doc.add_paragraph()
    birthday_year = data.get('birthday_year', '')
    birthday_month = data.get('birthday_month', '')
    birthday_day = data.get('birthday_day', '')
    run = pc_birthday.add_run(f"出生日期：{birthday_year}年 {birthday_month}月 {birthday_day}日")
    set_font_color_space(run, pc_birthday, 14)

    # 身份证号
    pc_idnumber = doc.add_paragraph()
    run = pc_idnumber.add_run("身份证号：" + data.get('id_number', ''))
    set_font_color_space(run, pc_idnumber, 14)

    # 户籍地
    pc_registered_address = doc.add_paragraph()
    run = pc_registered_address.add_run("户籍地：" + data.get('registered_address', ''))
    set_font_color_space(run, pc_registered_address, 14)

    # 现居住地
    pc_residence_address = doc.add_paragraph()
    run = pc_residence_address.add_run("现居住地：" + data.get('residence_address', ''))
    set_font_color_space(run, pc_residence_address, 14)

    # 联系电话
    pc_phone = doc.add_paragraph()
    run = pc_phone.add_run("联系电话：" + data.get('claimant_phone', ''))
    set_font_color_space(run, pc_phone, 14)

    # 空行
    doc.add_paragraph()

    # 被申请人
    p_respondent = doc.add_paragraph()
    run = p_respondent.add_run("被申请人：")
    run.bold = True
    set_font_color_space(run, p_respondent, 14)

    # 公司名称
    pr_company = doc.add_paragraph()
    run = pr_company.add_run("公司名称：" + data.get('respondent_company', ''))
    set_font_color_space(run, pr_company, 14)

    # 公司地址
    pr_company_address = doc.add_paragraph()
    run = pr_company_address.add_run("公司地址：" + data.get('respondent_address', ''))
    set_font_color_space(run, pr_company_address, 14)

    # 法定代表人、负责人
    pr_legal_representative = doc.add_paragraph()
    run = pr_legal_representative.add_run("法定代表人/负责人：" + data.get('legal_representative', ''))
    set_font_color_space(run, pr_legal_representative, 14)

    # 联系电话
    pr_phone = doc.add_paragraph()
    run = pr_phone.add_run("联系电话：" + data.get('respondent_phone', ''))
    set_font_color_space(run, pr_phone, 14)

    # 空行
    doc.add_paragraph()

    # 请求事项
    pr_requests = doc.add_paragraph()
    run = pr_requests.add_run("请求事项：")
    run.bold = True
    set_font_color_space(run, pr_requests, 14)

    # 裁令被申请人向申请人支付以下工伤待遇：
    pr_content = doc.add_paragraph()
    run = pr_content.add_run("裁令被申请人向申请人支付以下工伤待遇：")
    set_font_color_space(run, pr_content, 14)

    # 医疗费
    pc_medical_fees = doc.add_paragraph()
    amount_add_underline(pc_medical_fees, "1.医疗费", data.get('medical_fees', ''))

    # 页数计数器
    pageNum = 2

    # 伙食补助费
    if data.get('food_allowanc', ''):
        pc_food_allowance = doc.add_paragraph()
        fee_name = f"{pageNum}.伙食补助费"
        amount_add_underline(pc_food_allowance, fee_name, data.get('food_allowanc', ''))
        pageNum += 1

    # 交通费和食宿费（外地治疗）    
    if data.get('treatment_outtown', ''):
        pc_treatment_outtown = doc.add_paragraph()
        fee_name = f"{pageNum}.交通费和食宿费（外地治疗）"
        amount_add_underline(pc_treatment_outtown, fee_name, data.get('treatment_outtown', ''))
        pageNum += 1

    # 停工留薪期工资    
    if data.get('workstoppage_salary', ''):
        pc_workstoppage_salary = doc.add_paragraph()
        fee_name = f"{pageNum}.停工留薪期工资"
        amount_add_underline(pc_workstoppage_salary, fee_name, data.get('workstoppage_salary', ''))
        pageNum += 1

    # 护理费用    
    if data.get('nursing_fee', ''):     
        pc_nursing_fee = doc.add_paragraph()
        fee_name = f"{pageNum}.护理费用"
        amount_add_underline(pc_nursing_fee, fee_name, data.get('nursing_fee', ''))
        pageNum += 1

    # 工伤康复费用
    if data.get('rehabilitation_fee', ''):
        pc_rehabilitation_fee = doc.add_paragraph()
        fee_name = f"{pageNum}.工伤康复费用"
        amount_add_underline(pc_rehabilitation_fee, fee_name, data.get('rehabilitation_fee', ''))
        pageNum += 1

    # 一次性伤残补助金
    if data.get('onetime_disability', ''):
        pc_onetime_disability = doc.add_paragraph()
        fee_name = f"{pageNum}.一次性伤残补助金"
        amount_add_underline(pc_onetime_disability, fee_name, data.get('onetime_disability', ''))
        pageNum += 1

    # 一次性工伤医疗补助金
    if data.get('onetime_workinjury', ''):
        pc_onetime_workinjury = doc.add_paragraph()
        fee_name = f"{pageNum}.一次性工伤医疗补助金"
        amount_add_underline(pc_onetime_workinjury, fee_name, data.get('onetime_workinjury', ''))
        pageNum += 1

    # 一次性伤残就业补助金
    if data.get('onetime_disabilityemployment', ''):
        pc_onetime_disabilityemployment = doc.add_paragraph()
        fee_name = f"{pageNum}.一次性伤残就业补助金"
        amount_add_underline(pc_onetime_disabilityemployment, fee_name, data.get('onetime_disabilityemployment', ''))
        pageNum += 1

    # 每月伤残津贴
    if data.get('monthly_disabilityallowance', ''):
        pc_monthly_disabilityallowance = doc.add_paragraph()
        fee_name = f"{pageNum}.每月伤残津贴"
        amount_add_underline(pc_monthly_disabilityallowance, fee_name, data.get('monthly_disabilityallowance', ''))
        pageNum += 1

        # 每月伤残津贴起止时间
        pc_monthly_disabilityallowance_time = doc.add_paragraph()
        text = [
            '（自',
            'monthly_disabilityallowance_time_year',
            '年',
            'monthly_disabilityallowance_time_month',
            '月',
            'monthly_disabilityallowance_time_day',
            '日起支付至申请人退休日）'
        ]
        others_add_underline(pc_monthly_disabilityallowance_time, text, data)

    # 每月生活护理费
    if data.get('monthly-livingcare', ''):
        pc_monthly_livingcare = doc.add_paragraph()
        fee_name = f"{pageNum}.每月生活护理费"
        amount_add_underline(pc_monthly_livingcare, fee_name, data.get('monthly-livingcare', ''))
        pageNum += 1

        # 每月生活护理费起止时间
        pc_monthly_livingcare_time = doc.add_paragraph()
        text = [
            '（自',
            'monthly_livingcare_time_year',
            '年',
            'monthly_livingcare_time_month',
            '月',
            'monthly_livingcare_time_day',
            '日起支付至申请人退休日）'
        ]
        others_add_underline(pc_monthly_livingcare_time, text, data)

    # 合计费用
    total_fees = data.get('total_fees', '0')
    upper_total_fees = num_to_upper_main(total_fees) if total_fees else ''
    pc_total_fees = doc.add_paragraph()
    run = pc_total_fees.add_run("以上合计")
    set_font_color_space(run, pc_total_fees, 14)
    run = pc_total_fees.add_run("   " + total_fees + "   ")
    run.underline = True
    set_font_color_space(run, pc_total_fees, 14)
    run = pc_total_fees.add_run("元，合计大写金额")
    set_font_color_space(run, pc_total_fees, 14)
    run = pc_total_fees.add_run("   " + upper_total_fees + "   ")
    run.underline = True
    set_font_color_space(run, pc_total_fees, 14)
    run = pc_total_fees.add_run("圆。")
    set_font_color_space(run, pc_total_fees, 14)

    # 空行
    doc.add_paragraph()

    # 事实和理由
    pc_facts_reasons = doc.add_paragraph()
    run = pc_facts_reasons.add_run("事实和理由：")
    run.bold = True
    set_font_color_space(run, pc_facts_reasons, 14)

    # 事实和理由内容
    pc_facts_reasons_content = doc.add_paragraph()
    # 设置首行缩进为两个字符的宽度，这里假设使用14磅字号的字体，所以乘以2
    pc_facts_reasons_content.paragraph_format.first_line_indent = Pt(14 * 2)
    text = [
        "申请人自",
        'hiredate_year',
        "年",
        'hiredate_month',
        "月",
        'hiredate_day',
        "日入职，与被申请人建立劳动关系，工作岗位",
        'job',
        "，月工资标准",
        'monthly_salary_standard',
        "元。",
        'injury_certification_year',
        "年",
        'injury_certification_month',
        "月",
        'injury_certification_day',
        "日申请人受伤经过",
        'certification_organization',
        "，伤情诊断",
        'injury_condition',
        '，',
        'injury_certification_government_year',
        "年",
        'injury_certification_government_month',
        "月",
        'injury_certification_government_day',
        "日，经",
        'government_city',
        "市",
        'government_region',
        'selection_government_region',
        "人力资源和社会保障局认定为工伤，并经劳动能力鉴定为伤残等级",
        'disability_level',
        "级。现被申请人并未依法向申请人支付工伤相关待遇，严重损害了申请人的合法利益。"
    ]
    others_add_underline(pc_facts_reasons_content, text, data)

    # 依据相关法律法规，特向贵委提出仲裁申请，望裁如所请。
    p_content = doc.add_paragraph()
    p_content.paragraph_format.first_line_indent = Pt(14 * 2)
    run = p_content.add_run("依据相关法律法规，特向贵委提出仲裁申请，望裁如所请。")
    set_font_color_space(run, p_content, 14)

    # 此致
    p_content = doc.add_paragraph()
    p_content.paragraph_format.first_line_indent = Pt(14 * 2)
    run = p_content.add_run("此致")
    set_font_color_space(run, p_content, 14)

    # 市区劳动人事争议仲裁委员会
    p_area = doc.add_paragraph()
    p_area.paragraph_format.first_line_indent = Pt(14 * 2)
    text = [
        'application_city',
        "市",
        'application_region',
        'selection_application_region',
        "劳动人事争议仲裁委员会"
    ]
    others_add_underline(p_area, text, data)

    # 申请人签名
    p_sign = doc.add_paragraph()
    # 将缩进值设置为9.07厘米
    p_sign.paragraph_format.left_indent = Cm(9.07)
    run = p_sign.add_run("申请人：")
    set_font_color_space(run, p_sign, 14)

    # 申请日期
    p_date = doc.add_paragraph()
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p_date.add_run("      年      月      日")
    set_font_color_space(run, p_date, 14)

    return doc
